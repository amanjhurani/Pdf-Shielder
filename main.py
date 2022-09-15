from tkinter.ttk import Combobox
import openpyxl
import xlsxwriter 
from docx import Document
from docx.shared import RGBColor, Inches #Helps to specify font Color
from docx.shared import Inches
from docx.shared import Mm
from tkinter import *
import sys
import os
from libraries.applylogo import applyLogo
from default.default import COMPANY_DEFAULT, PASSWORD_DEFAULT

# Custom libraries
from libraries.pdfToImg import convertPdfToImage
from libraries.docTopdf import convertDocxToPdf
from libraries.imgTopdf import convertImgToPdf
from libraries.lockpdf import encryptPdf
from libraries.watermark import applyWatermark

wd = os.getcwd()

logo_mapper = {
    '0' : 'xyzCompany',
    '1' : 'abcCompany'
}

overlay_mapper = {
    '0' : 'exmrmp_o',
    '1' : 'extprp_o'
}


# Logo path

logo_path = wd + '\logo\{}.png'.format(logo_mapper[COMPANY_DEFAULT])
overlay_path = '\logo\{}.png'.format(overlay_mapper[COMPANY_DEFAULT])
only_logo_path = '\logo\{}.png'.format(logo_mapper[COMPANY_DEFAULT])


pwd = PASSWORD_DEFAULT


def onClick():
    root.withdraw()
    global logo_path
    global overlay_path
    global only_logo_path
    if (combobox.get() == "xyzCompany"):
        logo_path = wd + '\logo\{}.png'.format(logo_mapper['0'])
        overlay_path = '\logo\{}.png'.format(overlay_mapper['0'])
        only_logo_path = '\logo\{}.png'.format(logo_mapper['0'])
    elif (combobox.get() == "abcCompany"):
        print("------------->>>", combobox.get())
        logo_path = wd + '\logo\{}.png'.format(logo_mapper['1'])
        overlay_path = '\logo\{}.png'.format(overlay_mapper['1'])
        only_logo_path = '\logo\{}.png'.format(logo_mapper['1'])
    global pwd
    global HAVE_PDF
    global NEED_WATERMARK
    HAVE_PDF = int(havePdf.get())
    NEED_WATERMARK = int(needWatermark.get())
    pwd = password_entry.get()
    root.destroy()

def on_closing():
    sys.exit()

root = Tk()
root.title("EXCEL MODIFIER AND ENCRYPTER (LIFETIME ACCESS)")
root.geometry("460x480+100+100")


frame = Frame(root, bd=1, relief=None)
frame.pack(pady=5)

frame2 = Frame(root, bd=1, relief=None)
frame2.pack(pady=5)

rgb_label = Label(frame2, text="RGB COLOR CODE :", font=("Times New Roman", 12, "bold"))
rgb_label.grid(row=2, column=0)

havePdf = IntVar()
Checkbutton(frame2, text="PDF (sample.pdf) ✅/ DOC (sample_docx.docx) ❌?", variable=havePdf).grid(row=3, sticky=W)

needWatermark = IntVar(value=1)
Checkbutton(frame2, text="Need Watermark?", variable=needWatermark).grid(row=4, sticky=W)

password_label = Label(frame2, text="Pdf Password :", font=("Times New Roman", 12, "bold"))
password_label.grid(row=5, column=0)

password_entry = Entry(frame2, width=12, font=("Times New Roman", 12))
password_entry.grid(row=5, column=1, padx=5)
password_entry.insert(END, PASSWORD_DEFAULT)

combobox= Combobox(frame2,state= "readonly")
combobox['values']=('xyzCompany','abcCompany')
combobox.current(int(COMPANY_DEFAULT))
combobox.grid(row=6, column=1, columnspan=4)

copy = Button(frame2, text="SELECT", font=("Times New Roman", 12, "bold"),command = onClick)
copy.grid(row=7, columnspan=2, pady=7)

root.protocol("WM_DELETE_WINDOW", on_closing)
root.mainloop()





# #Coverting Docx to PDF
if (HAVE_PDF == 0):
    document = Document("sample_docx.docx")
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    header = section.header
    footer = section.footer
    header.add_paragraph("")
    paragraph = header.paragraphs[0]
    logo_run = paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1))
    document.save('sample_docx.docx')
    convertDocxToPdf('\sample_docx.docx', '\sample.pdf', wd)
pageCount = convertPdfToImage('sample.pdf', wd)
if NEED_WATERMARK == 1:
    applyWatermark(wd, overlay_path)
if HAVE_PDF == 1:
    applyLogo(wd, only_logo_path)
convertImgToPdf(wd, pageCount, 'output.pdf')
if (len(pwd) != 0):
    encryptPdf(wd,'output.pdf', pwd)
